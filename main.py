import os
import time
import glob
import pyautogui
from docx import Document
from docx.shared import Inches
import shutil

# -----------------------------------------------------
# SETTINGS
# -----------------------------------------------------
SCRIPTS_DIR = "scripts"
SCREENSHOT_DIR = "screenshots"
DOCX_OUTPUT = "SQL_Scripts.docx"
SCROLL_AMOUNT = -1500       # scroll down amount
MAX_SCROLLS = 50            # safety limit
VIEWER_CMD = 'code "{}"'    # VSCode open command
# VIEWER_CMD = 'notepad++ "{}"'  # If you prefer Notepad++

# Create screenshot directory
os.makedirs(SCREENSHOT_DIR, exist_ok=True)
shutil.rmtree(SCREENSHOT_DIR)

# -----------------------------------------------------
# TAKE MULTIPLE SCROLLING SCREENSHOTS
# -----------------------------------------------------
def capture_screens_for_file(sql_path, folder_name):
    target_dir = os.path.join(SCREENSHOT_DIR, folder_name)
    os.makedirs(target_dir, exist_ok=True)

    os.system(VIEWER_CMD.format(sql_path))
    time.sleep(2)  # wait for editor to open

    screenshots = []
    last_image = None

    for i in range(MAX_SCROLLS):
        img_path = os.path.join(target_dir, f"shot_{i}.png")

        screenshot = pyautogui.screenshot()
        screenshot.save(img_path)
        screenshots.append(img_path)

        # Detect end-of-file by comparing screenshots
        if last_image:
            if list(screenshot.getdata()) == last_image:
                print("Reached end of file.")
                break

        last_image = list(screenshot.getdata())

        pyautogui.scroll(SCROLL_AMOUNT)
        time.sleep(0.2)

    # Close VS Code window (Alt+F4)
    pyautogui.hotkey('alt', 'f4')
    time.sleep(0.3)

    return screenshots

# -----------------------------------------------------
# BUILD WORD DOCUMENT
# -----------------------------------------------------
def build_docx(all_scripts):
    doc = Document()
    doc.add_heading("SQL Scripts", level=1)

    for sql_path in all_scripts:
        file_name = os.path.basename(sql_path)
        section_name = file_name.replace(".sql", "")

        doc.add_heading(section_name, level=2)

        shot_folder = section_name

        pics = capture_screens_for_file(sql_path, shot_folder)

        for p in pics:
            doc.add_picture(p, width=Inches(6))
            doc.add_page_break()

    doc.save(DOCX_OUTPUT)
    print("DONE → Generated:", DOCX_OUTPUT)

# -----------------------------------------------------
# MAIN
# -----------------------------------------------------
if __name__ == "__main__":
    sql_files = glob.glob(os.path.join(SCRIPTS_DIR, "*.sql"))

    if not sql_files:
        print("No SQL files found in /scripts folder.")
        exit()

    print("Generating document…")
    build_docx(sql_files)
